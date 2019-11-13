from docx import Document 

document = Document()

#Create 'Header' object and add header template
header = document.sections[0] .header
header.add_paragraph('ENTER CUSTOM HEADER')


with open('ENTER FILEPATH OF HOMEWORKNUMBER.TXT FILE', 'r') as f:
	a = f.readlines()

#Intialize and increment homework # in text file
hmwkNum = int(a[0])
hmwkNum = hmwkNum + 1

with open ('ENTER FILEPATH OF HOMEWORKNUMBER.TXT FILE', 'w') as f:
	f.write(str(hmwkNum))

#Create 'Heading' object then bold and center the heading
heading = document.add_paragraph('')
heading.add_run("Homework #{0}".format(hmwkNum)).bold = True
heading.alignment = 1

#Create first question by adding 'Paragraph' object
document.add_paragraph('1.')

#Initialize String with 'Homework #'
fileName = "Homework #{0} - YOUR NAME.docx".format(hmwkNum)

#Initialize String with file path where document will be saved
filePath = 'ENTER FILEPATH OF WHERE THE OUTPUTTED FILE WILL BE SAVED' + fileName

#Save document to the file path
document.save(filePath)


